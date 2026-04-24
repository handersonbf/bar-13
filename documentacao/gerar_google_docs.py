from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = BASE_DIR / "Bar13-documentacao-google-docs.docx"

DOC_FILES = [
    ("Apresentação", BASE_DIR / "README.md"),
    ("Visão Geral", BASE_DIR / "visao-geral.md"),
    ("Mapa de Navegação", BASE_DIR / "mapa-de-navegacao.md"),
    ("Telas", BASE_DIR / "telas.md"),
    ("Funcionalidades", BASE_DIR / "funcionalidades.md"),
    ("Fluxo Principal", BASE_DIR / "fluxo-principal.md"),
    ("Arquitetura e Dados", BASE_DIR / "arquitetura-e-dados.md"),
]


def strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    return text.strip()


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 24, RGBColor(26, 26, 26)),
        ("Heading 1", 18, RGBColor(30, 30, 30)),
        ("Heading 2", 14, RGBColor(50, 50, 50)),
        ("Heading 3", 11.5, RGBColor(70, 70, 70)),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run("Bar13 • Documentação para Google Docs • Página ")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(110, 110, 110)
    add_page_number(footer_p)


def add_cover(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bar13")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(35, 35, 35)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Documentação funcional e técnica\nversão preparada para Google Docs")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(95, 95, 95)

    document.add_paragraph("")

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_run = intro.add_run(
        "Documento consolidado a partir dos arquivos Markdown da pasta documentacao,\n"
        "organizado para leitura, revisão e edição dentro do Google Docs."
    )
    intro_run.font.name = "Arial"
    intro_run.font.size = Pt(10.5)
    intro_run.font.color.rgb = RGBColor(120, 120, 120)

    document.add_page_break()


def add_index(document: Document) -> None:
    document.add_paragraph("Sumário", style="Heading 1")
    for title, _ in DOC_FILES:
        document.add_paragraph(title, style="List Bullet")
    document.add_page_break()


def add_markdown_file(document: Document, section_title: str, file_path: Path) -> None:
    document.add_paragraph(section_title, style="Heading 1")
    lines = file_path.read_text(encoding="utf-8").splitlines()

    in_code_block = False
    code_lines: list[str] = []

    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block and code_lines:
                document.add_paragraph("Bloco técnico:", style="Heading 3")
                for code_line in code_lines:
                    code_para = document.add_paragraph()
                    code_para.paragraph_format.left_indent = Inches(0.3)
                    code_run = code_para.add_run(code_line)
                    code_run.font.name = "Courier New"
                    code_run.font.size = Pt(9.5)
                    code_run.font.color.rgb = RGBColor(80, 80, 80)
                code_lines = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            continue

        if stripped.startswith("# "):
            continue

        if stripped.startswith("## "):
            document.add_paragraph(strip_inline_markdown(stripped[3:]), style="Heading 2")
            continue

        if stripped.startswith("### "):
            document.add_paragraph(strip_inline_markdown(stripped[4:]), style="Heading 3")
            continue

        if re.match(r"^\d+\.\s", stripped):
            document.add_paragraph(strip_inline_markdown(re.sub(r"^\d+\.\s+", "", stripped)), style="List Number")
            continue

        if stripped.startswith("- "):
            document.add_paragraph(strip_inline_markdown(stripped[2:]), style="List Bullet")
            continue

        para = document.add_paragraph(strip_inline_markdown(stripped))
        para.paragraph_format.space_after = Pt(6)


def build_document() -> Path:
    document = Document()
    configure_document(document)
    add_cover(document)
    add_index(document)

    for index, (title, path) in enumerate(DOC_FILES):
        add_markdown_file(document, title, path)
        if index != len(DOC_FILES) - 1:
            document.add_page_break()

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
